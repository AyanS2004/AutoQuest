import os
import re
import hashlib
from typing import List, Dict, Any
from datetime import datetime
import pandas as pd
from pathlib import Path
import logging
from io import StringIO

from PyPDF2 import PdfReader
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Enhanced document processor with support for multiple formats and intelligent chunking."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
    def process_excel(self, file_path: str) -> List[Dict[str, Any]]:
        """Process Excel files with enhanced metadata extraction."""
        try:
            df = pd.read_excel(file_path)
            df.fillna("", inplace=True)
            
            # Extract metadata
            metadata = {
                "file_path": file_path,
                "file_type": "excel",
                "rows": len(df),
                "columns": len(df.columns),
                "column_names": list(df.columns),
                "processed_at": datetime.now().isoformat()
            }
            
            chunks = []
            for idx, row in df.iterrows():
                # Create structured text from row
                row_text = "\n".join([f"{col}: {row[col]}" for col in df.columns if str(row[col]).strip()])
                
                if row_text.strip():
                    chunk = {
                        "text": row_text,
                        "metadata": {
                            **metadata,
                            "row_index": idx,
                            "chunk_id": f"excel_{idx}_{hashlib.md5(row_text.encode()).hexdigest()[:8]}"
                        }
                    }
                    chunks.append(chunk)
            
            logger.info(f"Processed Excel file: {len(chunks)} chunks created")
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing Excel file {file_path}: {str(e)}")
            raise
    
    def process_csv(self, file_path: str) -> List[Dict[str, Any]]:
        """Process CSV files with intelligent parsing."""
        try:
            df = pd.read_csv(file_path)
            df.fillna("", inplace=True)
            
            metadata = {
                "file_path": file_path,
                "file_type": "csv",
                "rows": len(df),
                "columns": len(df.columns),
                "column_names": list(df.columns),
                "processed_at": datetime.now().isoformat()
            }
            
            chunks = []
            for idx, row in df.iterrows():
                row_text = "\n".join([f"{col}: {row[col]}" for col in df.columns if str(row[col]).strip()])
                
                if row_text.strip():
                    chunk = {
                        "text": row_text,
                        "metadata": {
                            **metadata,
                            "row_index": idx,
                            "chunk_id": f"csv_{idx}_{hashlib.md5(row_text.encode()).hexdigest()[:8]}"
                        }
                    }
                    chunks.append(chunk)
            
            logger.info(f"Processed CSV file: {len(chunks)} chunks created")
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing CSV file {file_path}: {str(e)}")
            raise
    
    def process_text(self, file_path: str) -> List[Dict[str, Any]]:
        """Process text files with intelligent chunking."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            metadata = {
                "file_path": file_path,
                "file_type": "text",
                "file_size": len(content),
                "processed_at": datetime.now().isoformat()
            }
            
            # Split into paragraphs first
            paragraphs = re.split(r'\n\s*\n', content)
            
            chunks = []
            for i, paragraph in enumerate(paragraphs):
                if len(paragraph.strip()) < 50:  # Skip very short paragraphs
                    continue
                    
                # Further split long paragraphs
                if len(paragraph) > self.chunk_size:
                    sub_chunks = self._split_text(paragraph)
                    for j, sub_chunk in enumerate(sub_chunks):
                        chunk = {
                            "text": sub_chunk,
                            "metadata": {
                                **metadata,
                                "paragraph_index": i,
                                "sub_chunk_index": j,
                                "chunk_id": f"text_{i}_{j}_{hashlib.md5(sub_chunk.encode()).hexdigest()[:8]}"
                            }
                        }
                        chunks.append(chunk)
                else:
                    chunk = {
                        "text": paragraph,
                        "metadata": {
                            **metadata,
                            "paragraph_index": i,
                            "chunk_id": f"text_{i}_{hashlib.md5(paragraph.encode()).hexdigest()[:8]}"
                        }
                    }
                    chunks.append(chunk)
            
            logger.info(f"Processed text file: {len(chunks)} chunks created")
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {str(e)}")
            raise
    
    def _split_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks."""
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # Try to break at sentence boundaries
            if end < len(text):
                # Look for sentence endings
                for i in range(end, max(start + self.chunk_size - 100, start), -1):
                    if text[i] in '.!?':
                        end = i + 1
                        break
                else:
                    # Look for word boundaries
                    for i in range(end, max(start + self.chunk_size - 50, start), -1):
                        if text[i] == ' ':
                            end = i
                            break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - self.chunk_overlap
            if start >= len(text):
                break
        
        return chunks
    
    def process_file(self, file_path: str) -> List[Dict[str, Any]]:
        """Process any supported file type."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension in ['.xlsx', '.xls']:
            return self.process_excel(str(file_path))
        elif file_extension == '.csv':
            return self.process_csv(str(file_path))
        elif file_extension in ['.txt', '.md']:
            return self.process_text(str(file_path))
        elif file_extension == '.pdf':
            return self.process_pdf(str(file_path))
        elif file_extension == '.docx':
            return self.process_docx(str(file_path))
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def extract_keywords(self, text: str, top_k: int = 10) -> List[str]:
        """Extract key terms from text using simple frequency analysis."""
        # Remove common words and punctuation
        words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
        
        # Simple frequency analysis
        word_freq = {}
        for word in words:
            if word not in ['the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by']:
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # Return top k words
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        return [word for word, freq in sorted_words[:top_k]]
    
    def get_document_summary(self, chunks: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Generate a summary of the processed document."""
        total_chunks = len(chunks)
        total_text_length = sum(len(chunk['text']) for chunk in chunks)
        
        # Extract common keywords across all chunks
        all_text = " ".join(chunk['text'] for chunk in chunks)
        keywords = self.extract_keywords(all_text, top_k=15)
        
        # Get file metadata from first chunk
        file_metadata = chunks[0]['metadata'] if chunks else {}
        
        return {
            "total_chunks": total_chunks,
            "total_text_length": total_text_length,
            "average_chunk_length": total_text_length / total_chunks if total_chunks > 0 else 0,
            "keywords": keywords,
            "file_metadata": file_metadata,
            "processed_at": datetime.now().isoformat()
        } 

    # New: PDF processing
    def process_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """Process PDF files using PyPDF2 with page-level extraction and chunking."""
        try:
            reader = PdfReader(file_path)
            full_text_parts: List[str] = []
            page_map: List[int] = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                cleaned = self._clean_text(text)
                if cleaned:
                    full_text_parts.append(cleaned)
                    page_map.append(i + 1)
            full_text = "\n\n".join(full_text_parts)
            if not full_text.strip():
                return []
            # Chunk the combined text, but preserve approximate page indices via metadata
            chunks_text = self._split_text(full_text)
            chunks: List[Dict[str, Any]] = []
            for idx, chunk_text in enumerate(chunks_text):
                chunks.append({
                    "text": chunk_text,
                    "metadata": {
                        "file_path": file_path,
                        "file_type": "pdf",
                        "chunk_id": f"pdf_{idx}_{hashlib.md5(chunk_text.encode()).hexdigest()[:8]}",
                        "approx_page": None,
                        "processed_at": datetime.now().isoformat()
                    }
                })
            return chunks
        except Exception as e:
            logger.error(f"Error processing PDF file {file_path}: {str(e)}")
            raise

    # New: DOCX processing
    def process_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """Process DOCX files with paragraph and table text extraction."""
        try:
            doc = DocxDocument(file_path)
            parts: List[str] = []
            for para in doc.paragraphs:
                text = self._clean_text(para.text)
                if text:
                    parts.append(text)
            # Tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(self._clean_text(cell.text) for cell in row.cells if self._clean_text(cell.text))
                    if row_text:
                        parts.append(row_text)
            combined = "\n".join(parts)
            if not combined.strip():
                return []
            chunks_text = self._split_text(combined)
            chunks: List[Dict[str, Any]] = []
            for idx, chunk_text in enumerate(chunks_text):
                chunks.append({
                    "text": chunk_text,
                    "metadata": {
                        "file_path": file_path,
                        "file_type": "docx",
                        "chunk_id": f"docx_{idx}_{hashlib.md5(chunk_text.encode()).hexdigest()[:8]}",
                        "processed_at": datetime.now().isoformat()
                    }
                })
            return chunks
        except Exception as e:
            logger.error(f"Error processing DOCX file {file_path}: {str(e)}")
            raise

    # New: helper used by tests
    def _clean_text(self, text: str) -> str:
        """Normalize whitespace and trim."""
        if not text:
            return ""
        # Collapse multiple spaces and newlines
        text = re.sub(r"\s+", " ", str(text))
        return text.strip()

    # New: async API used in tests and higher-level flows
    async def process_document(self, file_path: str, document_type) -> List[Dict[str, Any]]:
        """Async wrapper to process a document by explicit type."""
        file_path = str(file_path)
        suffix = Path(file_path).suffix.lower()
        if document_type is None:
            # Fallback to extension-based
            return self.process_file(file_path)
        try:
            if suffix in [".xlsx", ".xls"]:
                return self.process_excel(file_path)
            if suffix == ".csv":
                return self.process_csv(file_path)
            if suffix in [".txt", ".md"]:
                return self.process_text(file_path)
            if suffix == ".pdf":
                return self.process_pdf(file_path)
            if suffix == ".docx":
                return self.process_docx(file_path)
            return self.process_file(file_path)
        except Exception:
            logger.exception(f"Failed to process document: {file_path}")
            raise